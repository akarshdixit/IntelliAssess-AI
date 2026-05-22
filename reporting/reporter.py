"""
reporting/reporter.py
======================
DOCX report generator for IntelliAssess AI — Phase 4-2.

Responsibility: consume structured intelligence and produce a professional
security assessment report in DOCX format.

  Input:  list[ParsedScanData]  — all parsed scanner outputs for the session
          EnrichedReport        — AI-enriched findings, executive summary, remediations
          Session               — metadata, context, targets, tools, processing log

  Output: security_report.docx — saved to the session's reports/ directory

Architecture position:
  EnrichedReport  ──→  DocxReporter.generate()  ──→  security_report.docx

Design principles:
  - Single public function: generate_docx() — stable API for Phase 4-3 integration
  - DocxReporter owns section generation — one private method per section
  - templates.py owns all formatting primitives — this module never touches XML
  - Graceful degradation: enriched=False triggers fallback prose/tables, never breaks
  - Never raises: errors are logged; the document is always saved

Report sections (in order):
  1. Cover Page              — session identity, risk posture, assessment metadata
  2. Executive Summary       — AI narrative or statistical fallback, key findings
  3. Assessment Scope        — targets, tools used, methodology description
  4. Technical Findings      — per-target grouped findings with severity, narrative,
                               compliance references, and contextual risk annotation
  5. TLS / HTTP / Infra      — structured observations from ssl_info and http_headers
  6. Positive Observations   — confirmed security controls, informational findings
  7. Remediation Roadmap     — time-horizoned action plan from AI remediation guidance
  8. Appendix                — session timeline, file inventory, enrichment status

Phase 5-1 enhancements (intelligence layer — display-only):
  - Technical Findings grouped by asset / target (report organisation only)
  - Compliance framework references appended to each finding block
    (PCI-DSS, CIS Controls, HIPAA, ISO 27001 — deterministic lookup)
  - Contextual severity adjustment annotated when context factors are active
    (exposure + environment + sector uplift — never mutates EnrichedReport)
  Both enhancements degrade gracefully when the intelligence/ package is absent.

Phase 4-3 integration note:
  session_manager.py will import and call:
    from reporting import generate_docx
    report_path = generate_docx(parsed_list, enriched_report, session, output_path)
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from reporting import templates as T
from utils.logger import get_logger

# Phase 5-1: lightweight intelligence layer (display-only, no upstream mutation)
try:
    from intelligence.compliance import get_compliance_refs
    from intelligence.risk_adjustment import adjust as _adjust_severity
    from intelligence.risk_adjustment import adjustment_delta, describe_adjustments
    _INTELLIGENCE_AVAILABLE = True
except ImportError:
    # Graceful degradation: if intelligence/ package is not yet on the path,
    # compliance refs and severity adjustment are silently suppressed.
    _INTELLIGENCE_AVAILABLE = False

    def get_compliance_refs(_ft: str) -> dict:           # type: ignore[misc]
        return {}

    def _adjust_severity(sev: str, _ctx: dict) -> str:  # type: ignore[misc]
        return sev

    def adjustment_delta(_sev: str, _ctx: dict) -> int: # type: ignore[misc]
        return 0

    def describe_adjustments(_ctx: dict) -> list:        # type: ignore[misc]
        return []

log = get_logger(__name__)

# Default output filename placed in the session's reports/ subdirectory.
REPORT_FILENAME = "security_report.docx"


# ---------------------------------------------------------------------------
# Public API — stable interface for Phase 4-3 session_manager integration
# ---------------------------------------------------------------------------

def generate_docx(
    parsed_data_list: list,          # list[ParsedScanData] — avoids circular import
    enriched:         object,        # EnrichedReport
    session:          object,        # Session
    output_path:      Optional[Path] = None,
) -> Path:
    """
    Generate a professional DOCX security assessment report.

    This is the single stable entry point for Phase 4-3.
    Internally constructs a DocxReporter, builds all sections, and saves.

    Args:
        parsed_data_list: All ParsedScanData objects for this session.
        enriched:         EnrichedReport from ai.analyzer.run().
        session:          Session dataclass with all metadata.
        output_path:      Override save path. When None, the document is saved
                          to the current working directory as REPORT_FILENAME.
                          Phase 4-3 will pass the session's reports/ directory.

    Returns:
        Path to the saved DOCX file.

    Contract:
        Never raises. On unexpected failure, logs the error and returns output_path
        (or a fallback path) so the caller can continue without crashing.
    """
    log.info(
        "Reporter: starting DOCX generation for session %s",
        getattr(session, 'session_id', 'UNKNOWN'),
    )

    try:
        reporter = DocxReporter(parsed_data_list, enriched, session)
        path     = reporter.generate(output_path)
        log.info("Reporter: DOCX saved → %s", path)
        return path

    except Exception as exc:
        log.error(
            "Reporter: unexpected error during DOCX generation: %s",
            exc,
            exc_info=True,
        )
        fallback = output_path or Path(REPORT_FILENAME)
        return fallback


# ---------------------------------------------------------------------------
# DocxReporter — section assembly
# ---------------------------------------------------------------------------

class DocxReporter:
    """
    Assembles a complete DOCX security assessment report from structured data.

    Each private method corresponds to one report section. Section order is
    controlled by generate(). All formatting delegates to templates.py helpers.

    Instance attributes:
        parsed_data_list: list[ParsedScanData]
        enriched:         EnrichedReport
        session:          Session
        doc:              python-docx Document (mutated throughout)
    """

    def __init__(
        self,
        parsed_data_list: list,
        enriched:         object,
        session:          object,
    ) -> None:
        self.parsed_data_list = parsed_data_list
        self.enriched         = enriched
        self.session          = session
        self.doc              = Document()
        T.configure_document_styles(self.doc)

    def generate(self, output_path: Optional[Path] = None) -> Path:
        """
        Build all sections in order and save the document.

        Returns the Path where the document was saved.
        """
        self._build_cover_page()
        self.doc.add_page_break()

        self._build_executive_summary()
        self.doc.add_page_break()

        self._build_assessment_scope()
        self.doc.add_page_break()

        self._build_technical_findings()
        self.doc.add_page_break()

        self._build_tls_http_observations()
        self.doc.add_page_break()

        self._build_positive_observations()
        self.doc.add_page_break()

        self._build_remediation_roadmap()
        self.doc.add_page_break()

        self._build_appendix()

        path = output_path or Path(REPORT_FILENAME)
        self.doc.save(str(path))
        return path

    # =========================================================================
    # Section 1 — Cover Page
    # =========================================================================

    def _build_cover_page(self) -> None:
        """
        Build a professional cover page.

        Layout:
          - Vertical spacer
          - "SECURITY ASSESSMENT REPORT" — large, navy, centered
          - Horizontal rule
          - Client / session label — large, centered
          - Assessment date — smaller, grey
          - Second horizontal rule
          - Risk posture badge table — color-coded severity
          - Assessment metadata table (session ID, exposure, environment, sector)
          - Confidentiality notice — small italic, centered
        """
        doc  = self.doc
        sess = self.session
        ctx  = getattr(sess, 'context', {}) or {}
        exec_sum = self.enriched.executive_summary

        assessment_date = _format_date(getattr(sess, 'created_at', ''))

        # ── Vertical padding before title ──────────────────────────────────
        T.add_spacer(doc, pt=48)

        # ── Report title ───────────────────────────────────────────────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run  = title_para.add_run("SECURITY ASSESSMENT REPORT")
        title_run.font.name = "Arial"
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        r, g, b = T.hex_to_rgb(T.COL_NAVY_HEX)
        title_run.font.color.rgb = RGBColor(r, g, b)

        T.add_horizontal_rule(doc)
        T.add_spacer(doc, pt=14)

        # ── Client / session label ─────────────────────────────────────────
        client_label = (
            ctx.get("company_name")
            or getattr(sess, 'client_label', None)
            or getattr(sess, 'session_id', 'Assessment')
        )
        label_para = doc.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.add_run(client_label)
        label_run.font.name = "Arial"
        label_run.font.size = Pt(18)
        label_run.font.bold = True

        T.add_spacer(doc, pt=6)

        # ── Assessment date ────────────────────────────────────────────────
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run  = date_para.add_run(f"Assessment Date: {assessment_date}")
        date_run.font.name = "Arial"
        date_run.font.size = Pt(11)
        r, g, b = T.hex_to_rgb(T.COL_GREY_HEX)
        date_run.font.color.rgb = RGBColor(r, g, b)

        T.add_spacer(doc, pt=20)
        T.add_horizontal_rule(doc)
        T.add_spacer(doc, pt=14)

        # ── Risk posture badge ─────────────────────────────────────────────
        risk_posture = getattr(exec_sum, 'risk_posture', 'UNKNOWN') or 'UNKNOWN'
        bg_hex, fg_hex = T.get_severity_colors(risk_posture)

        # 3-column table: left spacer | badge | right spacer
        risk_table = doc.add_table(rows=1, cols=3)
        risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        risk_table.cell(0, 0).width = Inches(2.5)
        risk_table.cell(0, 2).width = Inches(2.5)

        badge_cell = risk_table.cell(0, 1)
        badge_cell.width = Inches(2.5)
        T.shade_cell(badge_cell, bg_hex)
        T.set_cell_text(
            badge_cell,
            f"Risk Posture: {risk_posture}",
            bold=True,
            color_hex=fg_hex,
            font_size=12,
            center=True,
        )

        T.add_spacer(doc, pt=20)

        # ── Assessment metadata table ──────────────────────────────────────
        meta_rows = [
            ("Session ID",       getattr(sess, 'session_id', '—')),
            ("Assessment Type",  "External Non-Intrusive Security Assessment"),
            ("Exposure",         _ctx_value(ctx, "exposure")),
            ("Environment",      _ctx_value(ctx, "environment")),
            ("Sector",           _ctx_value(ctx, "sector")),
        ]
        if ctx.get("asset_owner"):
            meta_rows.append(("Asset Owner", ctx["asset_owner"]))
        if ctx.get("scope_notes"):
            meta_rows.append(("Scope",       ctx["scope_notes"]))

        meta_table = doc.add_table(rows=len(meta_rows), cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        T.apply_table_style(meta_table)

        for i, (label, value) in enumerate(meta_rows):
            lc = meta_table.cell(i, 0)
            vc = meta_table.cell(i, 1)
            lc.width = Inches(2.2)
            vc.width = Inches(4.5)
            T.shade_cell(lc, T.COL_LIGHT_BLUE)
            T.set_cell_text(lc, label, bold=True, color_hex=T.COL_NAVY_HEX, font_size=10)
            T.set_cell_text(vc, value, font_size=10)

        T.add_spacer(doc, pt=36)

        # ── Confidentiality notice ─────────────────────────────────────────
        notice_para = doc.add_paragraph()
        notice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notice_run  = notice_para.add_run(
            "CONFIDENTIAL  —  This report is intended solely for the named recipient. "
            "Unauthorized distribution is strictly prohibited."
        )
        notice_run.font.name   = "Arial"
        notice_run.font.size   = Pt(9)
        notice_run.font.italic = True
        r, g, b = T.hex_to_rgb("888888")
        notice_run.font.color.rgb = RGBColor(r, g, b)

    # =========================================================================
    # Section 2 — Executive Summary
    # =========================================================================

    def _build_executive_summary(self) -> None:
        """
        Build the Executive Summary section.

        When AIExecutiveSummary.enriched=True:
          - Renders overview_paragraph as the opening narrative
          - Lists key_findings as bullets
          - Includes priority_recommendation

        When enriched=False (Gemini unavailable):
          - Generates a data-driven overview from session metrics
          - Severity count table replaces narrative

        Always renders a severity count table for at-a-glance metrics.
        """
        doc      = self.doc
        exec_sum = self.enriched.executive_summary

        doc.add_heading("Executive Summary", level=1)
        T.add_horizontal_rule(doc)

        # ── Opening narrative ──────────────────────────────────────────────
        if getattr(exec_sum, 'enriched', False) and exec_sum.overview_paragraph:
            T.add_body_paragraph(doc, exec_sum.overview_paragraph)
        else:
            targets      = self.enriched.primary_targets or getattr(self.session, 'targets', [])
            finding_count = self.enriched.finding_count
            tools        = getattr(self.session, 'tools_detected', [])
            tool_str     = ", ".join(tools) if tools else "the available scan tools"

            T.add_body_paragraph(
                doc,
                f"This security assessment was conducted against {len(targets)} target(s) "
                f"using {tool_str}. A total of {finding_count} security finding(s) were "
                "identified during the assessment period. The findings are detailed in the "
                "Technical Findings section of this report."
            )

        T.add_spacer(doc)

        # ── Key findings ───────────────────────────────────────────────────
        key_findings = getattr(exec_sum, 'key_findings', []) or []
        if key_findings:
            doc.add_heading("Key Findings", level=2)
            for finding in key_findings:
                T.add_bullet_paragraph(doc, finding)
            T.add_spacer(doc)

        # ── Severity summary table — always rendered ───────────────────────
        severity_counts = self._compute_severity_counts()
        if severity_counts:
            doc.add_heading("Finding Severity Summary", level=2)
            self._build_severity_summary_table(severity_counts)
            T.add_spacer(doc)

        # ── Priority recommendation ────────────────────────────────────────
        priority_rec = getattr(exec_sum, 'priority_recommendation', '') or ''
        if priority_rec:
            doc.add_heading("Priority Recommendation", level=2)
            T.add_body_paragraph(doc, priority_rec)

    # =========================================================================
    # Section 3 — Assessment Scope
    # =========================================================================

    def _build_assessment_scope(self) -> None:
        """
        Build the Assessment Scope & Methodology section.

        Includes:
          - Assessed targets table (target index, hostname/IP)
          - Tools & techniques table (tool name, purpose description)
          - Methodology narrative
          - Optional scope notes from session context
        """
        doc  = self.doc
        sess = self.session
        ctx  = getattr(sess, 'context', {}) or {}

        doc.add_heading("Assessment Scope & Methodology", level=1)
        T.add_horizontal_rule(doc)

        T.add_body_paragraph(
            doc,
            "This section documents the scope of the security assessment, including "
            "assessed targets, tools used, and methodology applied."
        )
        T.add_spacer(doc)

        # ── Assessed targets ───────────────────────────────────────────────
        doc.add_heading("Assessed Targets", level=2)
        targets = self.enriched.primary_targets or getattr(sess, 'targets', [])

        if targets:
            tgt_table = doc.add_table(rows=len(targets) + 1, cols=2)
            T.apply_table_style(tgt_table)

            # Header row
            T.shade_cell(tgt_table.cell(0, 0), T.COL_NAVY_HEX)
            T.shade_cell(tgt_table.cell(0, 1), T.COL_NAVY_HEX)
            T.set_cell_text(tgt_table.cell(0, 0), "#",      bold=True, color_hex=T.COL_WHITE, font_size=10)
            T.set_cell_text(tgt_table.cell(0, 1), "Target", bold=True, color_hex=T.COL_WHITE, font_size=10)
            tgt_table.cell(0, 0).width = Inches(0.7)
            tgt_table.cell(0, 1).width = Inches(6.8)

            for i, tgt in enumerate(targets, start=1):
                row_bg = T.COL_LIGHT_GREY if i % 2 == 0 else T.COL_WHITE
                T.shade_cell(tgt_table.cell(i, 0), row_bg)
                T.shade_cell(tgt_table.cell(i, 1), row_bg)
                T.set_cell_text(tgt_table.cell(i, 0), str(i))
                T.set_cell_text(tgt_table.cell(i, 1), tgt)
        else:
            T.add_body_paragraph(doc, "No targets were recorded for this session.")

        T.add_spacer(doc)

        # ── Tools & techniques ─────────────────────────────────────────────
        doc.add_heading("Tools & Techniques", level=2)
        tools = getattr(sess, 'tools_detected', []) or []

        _TOOL_DESCRIPTIONS = {
            "NMAP":      "Network port scanning and service version detection.",
            "HTTPX":     "HTTP response analysis, header inspection, and technology fingerprinting.",
            "SSLSCAN":   "SSL/TLS protocol enumeration, cipher suite analysis, and certificate inspection.",
            "SUBFINDER": "Passive subdomain enumeration via public DNS and certificate transparency logs.",
        }

        if tools:
            tool_table = doc.add_table(rows=len(tools) + 1, cols=2)
            T.apply_table_style(tool_table)

            T.shade_cell(tool_table.cell(0, 0), T.COL_NAVY_HEX)
            T.shade_cell(tool_table.cell(0, 1), T.COL_NAVY_HEX)
            T.set_cell_text(tool_table.cell(0, 0), "Tool",    bold=True, color_hex=T.COL_WHITE, font_size=10)
            T.set_cell_text(tool_table.cell(0, 1), "Purpose", bold=True, color_hex=T.COL_WHITE, font_size=10)
            tool_table.cell(0, 0).width = Inches(1.8)
            tool_table.cell(0, 1).width = Inches(5.7)

            for i, tool_name in enumerate(tools, start=1):
                row_bg = T.COL_LIGHT_GREY if i % 2 == 0 else T.COL_WHITE
                desc   = _TOOL_DESCRIPTIONS.get(tool_name.upper(), "Security assessment tool.")
                T.shade_cell(tool_table.cell(i, 0), row_bg)
                T.shade_cell(tool_table.cell(i, 1), row_bg)
                T.set_cell_text(tool_table.cell(i, 0), tool_name)
                T.set_cell_text(tool_table.cell(i, 1), desc)
        else:
            T.add_body_paragraph(doc, "No tools were recorded for this session.")

        T.add_spacer(doc)

        # ── Methodology narrative ──────────────────────────────────────────
        doc.add_heading("Methodology", level=2)
        T.add_body_paragraph(
            doc,
            "This assessment was conducted using non-intrusive external reconnaissance "
            "techniques. All scanning was limited to the agreed scope and did not include "
            "active exploitation of any identified vulnerability. Findings are derived from "
            "observed service behaviour, protocol-level analysis, and publicly observable "
            "infrastructure characteristics."
        )
        T.add_body_paragraph(
            doc,
            "The assessment covers the following domains: network port exposure, SSL/TLS "
            "configuration quality, HTTP security header presence, service version disclosure, "
            "and infrastructure technology fingerprinting."
        )

        if ctx.get("scope_notes"):
            T.add_spacer(doc)
            T.add_label_paragraph(doc, "Scope Notes")
            T.add_body_paragraph(doc, ctx["scope_notes"])

        if ctx.get("infra_notes"):
            T.add_spacer(doc)
            T.add_label_paragraph(doc, "Infrastructure Notes")
            T.add_body_paragraph(doc, ctx["infra_notes"])

    # =========================================================================
    # Section 4 — Technical Findings
    # =========================================================================

    def _build_technical_findings(self) -> None:
        """
        Build the Technical Findings section.

        Phase 5-1: Findings are grouped by asset / target. Each target receives
        a dedicated sub-heading, followed by all findings for that target in
        severity order (as delivered by the enrichment layer).

        This is REPORT-LEVEL grouping only:
          - No graph logic, no attack chains, no finding synthesis.
          - The grouping key is finding.target (a hostname or IP string).
          - Targets are rendered in alphabetical order for determinism.
          - A "Multiple / Unknown Target" bucket catches findings with no target.

        Each finding block includes:
          - Context-adjusted severity badge (with annotation if severity changed)
          - Finding title, port, observation, business impact
          - Compliance framework references (PCI-DSS, CIS, HIPAA, ISO 27001)
          - Remediation guidance with technical commands
        """
        doc      = self.doc
        enriched = self.enriched
        ctx      = getattr(self.session, 'context', {}) or {}

        doc.add_heading("Technical Findings", level=1)
        T.add_horizontal_rule(doc)

        # Filter out pure INFO findings — those go to Positive Observations.
        security_findings = [
            f for f in enriched.finding_summaries
            if (f.severity_label or "INFO").upper() != "INFO"
        ]

        if not security_findings:
            T.add_body_paragraph(
                doc,
                "No technical security findings were identified during this assessment, "
                "or all findings are informational. Refer to the Positive Security "
                "Observations section for confirmed security controls."
            )
            return

        T.add_body_paragraph(
            doc,
            f"This section presents {len(security_findings)} security finding(s) identified "
            "during the assessment, organised by assessed asset. Each finding includes an "
            "analyst narrative, business impact statement, compliance references, and "
            "remediation guidance."
        )
        T.add_spacer(doc)

        # ── Group findings by target ───────────────────────────────────────
        # Build ordered dict: target_name → [findings].
        # Findings within each target group retain their original severity
        # ordering from the enrichment layer.
        _UNKNOWN_TARGET = "Unattributed Findings"
        grouped: dict[str, list] = {}

        for f in security_findings:
            target_key = (getattr(f, 'target', '') or '').strip() or _UNKNOWN_TARGET
            grouped.setdefault(target_key, []).append(f)

        # Render each target group in alphabetical order.
        # Move _UNKNOWN_TARGET to the end if present.
        target_order = sorted(
            [k for k in grouped if k != _UNKNOWN_TARGET]
        )
        if _UNKNOWN_TARGET in grouped:
            target_order.append(_UNKNOWN_TARGET)

        # ── Per-target rendering ───────────────────────────────────────────
        global_idx = 1   # Sequential finding number across all targets

        for target_name in target_order:
            target_findings = grouped[target_name]

            # Target header block
            self._render_target_header(target_name, target_findings, ctx)
            T.add_spacer(doc, pt=4)

            # Individual finding blocks under this target
            for finding in target_findings:
                self._render_finding_block(global_idx, finding, ctx)
                T.add_spacer(doc, pt=10)
                global_idx += 1

    def _render_target_header(
        self,
        target_name: str,
        findings:    list,
        ctx:         dict,
    ) -> None:
        """
        Render a target asset section header.

        Shows:
          - Target heading (level 2)
          - Finding count summary for this target
          - Context-adjustment notice if any uplift factors are active

        This is purely presentational — no data is mutated.
        """
        doc = self.doc

        doc.add_heading(f"Asset: {target_name}", level=2)

        # Summary line: X findings across Y severity levels
        sev_counts: dict[str, int] = {}
        for f in findings:
            sev = (getattr(f, 'severity_label', 'INFO') or 'INFO').upper()
            sev_counts[sev] = sev_counts.get(sev, 0) + 1

        ordered_sevs = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"]
        sev_parts    = [
            f"{sev_counts[s]} {s}" for s in ordered_sevs if s in sev_counts
        ]
        sev_summary  = ",  ".join(sev_parts) if sev_parts else "—"

        T.add_body_paragraph(
            doc,
            f"Findings for this asset: {len(findings)}  ({sev_summary})"
        )

        # Context-adjustment notice (only when intelligence layer is available
        # and at least one uplift factor is active).
        adjustment_reasons = describe_adjustments(ctx)
        if adjustment_reasons:
            notice_parts = ";  ".join(adjustment_reasons)
            T.add_label_paragraph(doc, "Context-Adjusted Risk")
            T.add_body_paragraph(
                doc,
                f"Severity ratings for this assessment have been contextually adjusted. "
                f"Active factors: {notice_parts}. "
                f"Adjusted severities are shown in finding headers below."
            )

        T.add_spacer(doc, pt=6)

    def _render_finding_block(self, idx: int, finding, ctx: dict | None = None) -> None:
        """
        Render one AIFindingSummary as a rich structured block.

        Phase 5-1 additions:
          - Context-adjusted severity: if uplift factors are active, the badge
            shows the adjusted severity. A "(context-adjusted)" annotation is
            appended to the finding title when the severity was raised.
          - Compliance references: a "Compliance References" row is appended
            after the remediation block for all finding types with mappings.
            Suppressed silently when no mapping exists.

        Args:
            idx:     Sequential finding number (1-based, global across targets).
            finding: AIFindingSummary dataclass instance.
            ctx:     Session context dict (exposure, environment, sector).
                     May be None — all adjustments degrade gracefully.
        """
        doc = self.doc
        ctx = ctx or getattr(self.session, 'context', {}) or {}

        # ── Raw attributes from finding ────────────────────────────────────
        finding_type    = getattr(finding, 'finding_type', 'UNKNOWN') or 'UNKNOWN'
        target          = getattr(finding, 'target', '') or ''
        port            = getattr(finding, 'port', None)
        is_enriched     = getattr(finding, 'enriched', False)
        narrative       = getattr(finding, 'analyst_narrative', '') or ''
        business_impact = getattr(finding, 'business_impact', '') or ''
        raw_detail      = getattr(finding, 'raw_finding_detail', '') or ''
        raw_severity    = (getattr(finding, 'severity_label', 'INFO') or 'INFO').upper()

        # ── Phase 5-1: Contextual severity adjustment ─────────────────────
        # Compute adjusted severity at render time.
        # Never mutates the EnrichedReport — this is display-only arithmetic.
        adjusted_severity = _adjust_severity(raw_severity, ctx)
        severity_was_raised = (
            adjustment_delta(raw_severity, ctx) > 0
        )
        # Use the adjusted severity for the visual badge.
        display_severity = adjusted_severity
        bg_hex, fg_hex   = T.get_severity_colors(display_severity)

        # ── Severity badge + title row ─────────────────────────────────────
        hdr = doc.add_table(rows=1, cols=2)
        T.apply_table_style(hdr)
        hdr.cell(0, 0).width = Inches(1.1)
        hdr.cell(0, 1).width = Inches(6.4)

        badge_cell = hdr.cell(0, 0)
        T.shade_cell(badge_cell, bg_hex)
        # Show adjusted severity in badge; add asterisk when raised.
        badge_label = display_severity + ("*" if severity_was_raised else "")
        T.set_cell_text(badge_cell, badge_label, bold=True, color_hex=fg_hex, font_size=10, center=True)

        title_cell = hdr.cell(0, 1)
        T.shade_cell(title_cell, T.COL_NAVY_HEX)
        title_text = f"Finding {idx}:  {finding_type.replace('_', ' ').title()}"
        if target:
            title_text += f"  —  {target}"
        if severity_was_raised:
            title_text += "  (context-adjusted)"
        T.set_cell_text(title_cell, title_text, bold=True, color_hex=T.COL_WHITE, font_size=11)

        T.add_spacer(doc, pt=4)

        # ── Severity annotation (when raised) ─────────────────────────────
        if severity_was_raised:
            T.add_label_paragraph(doc, "Severity Adjustment Note")
            T.add_body_paragraph(
                doc,
                f"Base severity: {raw_severity}  →  Context-adjusted: {adjusted_severity}. "
                f"Uplift factors: {';  '.join(describe_adjustments(ctx))}."
            )

        # ── Port / service context ─────────────────────────────────────────
        port_str = f"Port {port}" if port else "N/A"
        T.add_label_paragraph(doc, "Port / Service")
        T.add_body_paragraph(doc, port_str)

        # ── Observation ────────────────────────────────────────────────────
        T.add_label_paragraph(doc, "Observation")
        if is_enriched and narrative:
            T.add_body_paragraph(doc, narrative)
        else:
            T.add_body_paragraph(doc, raw_detail or "No detail available for this finding.")

        # ── Business impact ────────────────────────────────────────────────
        if is_enriched and business_impact:
            T.add_label_paragraph(doc, "Business Impact")
            T.add_body_paragraph(doc, business_impact)

        # ── Remediation guidance ───────────────────────────────────────────
        remediation = self.enriched.get_remediation(finding_type)
        if remediation and getattr(remediation, 'enriched', False):
            immediate   = getattr(remediation, 'immediate_actions', []) or []
            short_term  = getattr(remediation, 'short_term_actions', []) or []
            commands    = getattr(remediation, 'commands', []) or []
            references  = getattr(remediation, 'references', []) or []

            if immediate or short_term:
                T.add_label_paragraph(doc, "Recommendation")
                for action in immediate:
                    T.add_bullet_paragraph(doc, action)
                for action in short_term:
                    T.add_bullet_paragraph(doc, action)

            if commands:
                T.add_label_paragraph(doc, "Technical Commands")
                for cmd in commands:
                    T.add_code_paragraph(doc, cmd)

            if references:
                T.add_label_paragraph(doc, "References")
                for ref in references:
                    T.add_bullet_paragraph(doc, ref)
        else:
            # Minimal fallback recommendation
            T.add_label_paragraph(doc, "Recommendation")
            T.add_body_paragraph(
                doc,
                "Consult vendor documentation and relevant security standards to "
                "address this finding. Prioritize based on the severity rating above."
            )

        # ── Phase 5-1: Compliance references ──────────────────────────────
        # Deterministic lookup — no AI, no scoring, no reasoning.
        # Renders a two-column table: Framework | Control References
        # Suppressed silently when no mapping exists for this finding type.
        compliance_refs = get_compliance_refs(finding_type)
        if compliance_refs:
            T.add_label_paragraph(doc, "Compliance References")
            ref_rows = list(compliance_refs.items())

            comp_table = doc.add_table(rows=len(ref_rows) + 1, cols=2)
            T.apply_table_style(comp_table)

            # Header row
            T.shade_cell(comp_table.cell(0, 0), T.COL_NAVY_HEX)
            T.shade_cell(comp_table.cell(0, 1), T.COL_NAVY_HEX)
            T.set_cell_text(comp_table.cell(0, 0), "Framework",       bold=True, color_hex=T.COL_WHITE, font_size=9)
            T.set_cell_text(comp_table.cell(0, 1), "Control Reference", bold=True, color_hex=T.COL_WHITE, font_size=9)
            comp_table.cell(0, 0).width = Inches(1.8)
            comp_table.cell(0, 1).width = Inches(5.7)

            # Data rows
            for i, (framework, controls) in enumerate(ref_rows, start=1):
                row_bg = T.COL_LIGHT_GREY if i % 2 == 0 else T.COL_WHITE
                T.shade_cell(comp_table.cell(i, 0), row_bg)
                T.shade_cell(comp_table.cell(i, 1), row_bg)
                T.set_cell_text(comp_table.cell(i, 0), framework,              font_size=9)
                T.set_cell_text(comp_table.cell(i, 1), ",  ".join(controls),   font_size=9)

        T.add_horizontal_rule(doc)

    # =========================================================================
    # Section 5 — TLS / HTTP / Infrastructure Observations
    # =========================================================================

    def _build_tls_http_observations(self) -> None:
        """
        Build the TLS, HTTP, and Infrastructure Observations section.

        Draws exclusively from ParsedScanData (deterministic layer):
          - ssl_info   → TLS protocol support, cipher suites, certificate details
          - http_headers → security header presence/absence, technology disclosure
          - open_ports + scan_metadata → network exposure summary

        This section is deterministic — no AI enrichment is consumed here.
        AI-generated narratives for TLS/HTTP findings appear in Technical Findings.
        """
        doc = self.doc

        doc.add_heading("TLS / HTTP / Infrastructure Observations", level=1)
        T.add_horizontal_rule(doc)

        T.add_body_paragraph(
            doc,
            "The following observations are derived from direct analysis of TLS/SSL "
            "configuration, HTTP response behaviour, and network infrastructure "
            "characteristics identified during the assessment."
        )
        T.add_spacer(doc)

        rendered_any = False

        for parsed in self.parsed_data_list:
            ssl_info     = getattr(parsed, 'ssl_info', {}) or {}
            http_headers = getattr(parsed, 'http_headers', {}) or {}
            open_ports   = getattr(parsed, 'open_ports', []) or []

            if ssl_info:
                self._render_ssl_block(parsed)
                T.add_spacer(doc)
                rendered_any = True

            if http_headers:
                self._render_http_block(parsed)
                T.add_spacer(doc)
                rendered_any = True

            if open_ports:
                self._render_network_block(parsed)
                T.add_spacer(doc)
                rendered_any = True

        if not rendered_any:
            T.add_body_paragraph(
                doc,
                "No TLS, HTTP, or network infrastructure data was available for this session. "
                "Ensure that SSLScan and Httpx output files were included in the ingested file set."
            )

    def _render_ssl_block(self, parsed) -> None:
        """Render SSL/TLS observations for one ParsedScanData."""
        doc      = self.doc
        ssl_info = getattr(parsed, 'ssl_info', {}) or {}
        target   = getattr(parsed, 'primary_target', '') or 'Unknown Target'

        doc.add_heading(f"SSL/TLS Configuration — {target}", level=2)

        protocols       = ssl_info.get("supported_protocols", []) or []
        cipher_inventory= ssl_info.get("cipher_inventory", []) or []
        cert            = ssl_info.get("certificate", {}) or {}

        # ── Protocol support summary ───────────────────────────────────────
        # Protocol dicts are produced by the SSLScan parser. Accept either the
        # "name" or "protocol" key defensively so a parser dict-shape change can
        # never crash the mission-critical report. Skip entries with no label.
        def _proto_label(p: dict) -> str:
            return (p.get("name") or p.get("protocol") or "").strip()

        enabled_protos  = [lbl for p in protocols
                           if p.get("enabled") and (lbl := _proto_label(p))]
        disabled_protos = [lbl for p in protocols
                           if not p.get("enabled") and (lbl := _proto_label(p))]

        if enabled_protos or disabled_protos:
            T.add_label_paragraph(doc, "Protocol Support")
            for proto in enabled_protos:
                T.add_bullet_paragraph(doc, f"{proto}:  Enabled")
            if disabled_protos:
                T.add_bullet_paragraph(
                    doc,
                    f"Legacy protocols disabled: {', '.join(disabled_protos)}"
                )
            T.add_spacer(doc, pt=4)

        # ── Certificate details ────────────────────────────────────────────
        if cert:
            cert_rows = []
            if cert.get("subject"):           cert_rows.append(("Subject",          cert["subject"]))
            if cert.get("issuer"):            cert_rows.append(("Issuer",           cert["issuer"]))
            if cert.get("not_before"):        cert_rows.append(("Valid From",       cert["not_before"]))
            if cert.get("not_after"):         cert_rows.append(("Valid To",         cert["not_after"]))
            if cert.get("key_size"):          cert_rows.append(("RSA Key Size",     f"{cert['key_size']} bits"))
            if cert.get("sig_algorithm"):     cert_rows.append(("Signature Alg.",   cert["sig_algorithm"]))
            self_signed = cert.get("self_signed", False)
            cert_rows.append(("Self-Signed", "Yes — requires review" if self_signed else "No"))

            if cert_rows:
                T.add_label_paragraph(doc, "Certificate Details")
                ct = doc.add_table(rows=len(cert_rows), cols=2)
                T.apply_table_style(ct)
                for i, (k, v) in enumerate(cert_rows):
                    lc = ct.cell(i, 0)
                    vc = ct.cell(i, 1)
                    lc.width = Inches(2.0)
                    vc.width = Inches(5.5)
                    T.shade_cell(lc, T.COL_LIGHT_BLUE)
                    T.set_cell_text(lc, k, bold=True, color_hex=T.COL_NAVY_HEX, font_size=10)
                    T.set_cell_text(vc, str(v), font_size=10)
                T.add_spacer(doc, pt=4)

        # ── Accepted cipher suites ─────────────────────────────────────────
        # Identify weak ciphers by cross-referencing WEAK_CIPHER findings
        weak_cipher_names = self._collect_weak_cipher_names(target)

        if cipher_inventory:
            strong = [c["cipher"] for c in cipher_inventory if c["cipher"] not in weak_cipher_names]
            weak   = [c["cipher"] for c in cipher_inventory if c["cipher"] in weak_cipher_names]

            if strong:
                T.add_label_paragraph(doc, "Accepted Strong Cipher Suites")
                for cipher_name in strong[:12]:  # Cap display at 12 to avoid wall-of-text
                    T.add_bullet_paragraph(doc, cipher_name)
                if len(strong) > 12:
                    T.add_body_paragraph(
                        doc, f"  ... and {len(strong) - 12} additional cipher suite(s)."
                    )

            if weak:
                T.add_spacer(doc, pt=4)
                T.add_label_paragraph(doc, "Weak Cipher Suites Detected")
                for cipher_name in weak:
                    T.add_bullet_paragraph(doc, f"\u26a0  {cipher_name}")

    def _render_http_block(self, parsed) -> None:
        """Render HTTP security header observations for one ParsedScanData."""
        doc     = self.doc
        headers = getattr(parsed, 'http_headers', {}) or {}
        target  = getattr(parsed, 'primary_target', '') or 'Unknown Target'

        doc.add_heading(f"HTTP Security Headers — {target}", level=2)

        # Expected security headers: lowercase key → display name
        _SECURITY_HEADERS = {
            "strict-transport-security": "Strict-Transport-Security (HSTS)",
            "content-security-policy":   "Content-Security-Policy",
            "x-frame-options":           "X-Frame-Options",
            "x-content-type-options":    "X-Content-Type-Options",
            "referrer-policy":           "Referrer-Policy",
            "permissions-policy":        "Permissions-Policy",
        }

        headers_lc = {k.lower(): v for k, v in headers.items()}

        present = {
            display: headers_lc[key]
            for key, display in _SECURITY_HEADERS.items()
            if key in headers_lc
        }
        missing = {
            display
            for key, display in _SECURITY_HEADERS.items()
            if key not in headers_lc
        }

        if present:
            T.add_label_paragraph(doc, "Security Headers Present")
            for hdr_name, hdr_val in present.items():
                display_val = (hdr_val[:80] + "...") if len(str(hdr_val)) > 80 else str(hdr_val)
                T.add_bullet_paragraph(doc, f"{hdr_name}: {display_val}")
            T.add_spacer(doc, pt=4)

        if missing:
            T.add_label_paragraph(doc, "Security Headers Missing")
            for hdr_name in sorted(missing):
                T.add_bullet_paragraph(doc, hdr_name)
            T.add_spacer(doc, pt=4)

        # Server / technology disclosure
        server_val = headers_lc.get("server", "")
        if server_val:
            T.add_label_paragraph(doc, "Server Header (Technology Disclosure)")
            T.add_body_paragraph(doc, server_val)

        # X-Powered-By
        powered_by = headers_lc.get("x-powered-by", "")
        if powered_by:
            T.add_label_paragraph(doc, "X-Powered-By")
            T.add_body_paragraph(doc, powered_by)

    def _render_network_block(self, parsed) -> None:
        """Render open port / network exposure summary for one ParsedScanData."""
        doc        = self.doc
        open_ports = getattr(parsed, 'open_ports', []) or []
        target     = getattr(parsed, 'primary_target', '') or 'Unknown Target'
        scan_meta  = getattr(parsed, 'scan_metadata', {}) or {}
        assets     = getattr(parsed, 'assets', []) or []

        doc.add_heading(f"Network Exposure — {target}", level=2)

        port_count = scan_meta.get("port_count")
        T.add_body_paragraph(
            doc,
            f"Open ports identified: {len(open_ports)}"
            + (f" (of {port_count} scanned)" if port_count else "")
        )

        if open_ports:
            # Build port-service table from assets
            port_service_map: dict[int, str] = {}
            for asset in assets:
                for svc in getattr(asset, 'services', []):
                    if getattr(svc, 'state', '') == 'open':
                        port_service_map[svc.port] = (
                            f"{svc.service_name or 'unknown'}"
                            + (f" ({svc.version})" if getattr(svc, 'version', None) else "")
                        )

            T.add_label_paragraph(doc, "Open Ports")
            port_table = doc.add_table(rows=len(open_ports) + 1, cols=2)
            T.apply_table_style(port_table)
            T.shade_cell(port_table.cell(0, 0), T.COL_NAVY_HEX)
            T.shade_cell(port_table.cell(0, 1), T.COL_NAVY_HEX)
            T.set_cell_text(port_table.cell(0, 0), "Port",    bold=True, color_hex=T.COL_WHITE, font_size=10)
            T.set_cell_text(port_table.cell(0, 1), "Service", bold=True, color_hex=T.COL_WHITE, font_size=10)
            port_table.cell(0, 0).width = Inches(1.5)
            port_table.cell(0, 1).width = Inches(6.0)

            for i, port in enumerate(open_ports, start=1):
                row_bg  = T.COL_LIGHT_GREY if i % 2 == 0 else T.COL_WHITE
                svc_str = port_service_map.get(port, "—")
                T.shade_cell(port_table.cell(i, 0), row_bg)
                T.shade_cell(port_table.cell(i, 1), row_bg)
                T.set_cell_text(port_table.cell(i, 0), str(port))
                T.set_cell_text(port_table.cell(i, 1), svc_str)

    # =========================================================================
    # Section 6 — Positive Security Observations
    # =========================================================================

    def _build_positive_observations(self) -> None:
        """
        Build the Positive Security Observations section.

        Sources:
          - AIExecutiveSummary.positive_observations (AI-sourced controls list)
          - INFO-severity AIFindingSummary entries (TLS_ENABLED, etc.)

        Balanced reporting: confirms what is working, not just what is broken.
        """
        doc      = self.doc
        exec_sum = self.enriched.executive_summary

        doc.add_heading("Positive Security Observations", level=1)
        T.add_horizontal_rule(doc)

        T.add_body_paragraph(
            doc,
            "The following security controls were observed to be correctly implemented "
            "at the time of assessment. These represent areas of strength that should "
            "be maintained and continuously reviewed."
        )
        T.add_spacer(doc)

        positive_notes = getattr(exec_sum, 'positive_observations', []) or []

        # INFO-severity findings: TLS_ENABLED, correctly configured controls
        info_findings = [
            f for f in self.enriched.finding_summaries
            if (f.severity_label or "INFO").upper() == "INFO"
        ]

        if not positive_notes and not info_findings:
            T.add_body_paragraph(
                doc,
                "No positive security observations were confirmed during this assessment. "
                "This may indicate insufficient data from the scan set, or that confirmable "
                "positive controls were not within scope."
            )
            return

        # ── AI-sourced positive observations ──────────────────────────────
        if positive_notes:
            doc.add_heading("Confirmed Security Controls", level=2)
            for note in positive_notes:
                T.add_bullet_paragraph(doc, note)
            T.add_spacer(doc)

        # ── INFO findings as positive evidence ────────────────────────────
        if info_findings:
            doc.add_heading("Informational Observations", level=2)
            for finding in info_findings:
                detail = (
                    getattr(finding, 'raw_finding_detail', '') or
                    (finding.finding_type or '').replace('_', ' ').title()
                )
                T.add_bullet_paragraph(doc, detail)

    # =========================================================================
    # Section 7 — Remediation Roadmap
    # =========================================================================

    def _build_remediation_roadmap(self) -> None:
        """
        Build the Remediation Roadmap section.

        Organizes AIRemediation entries across three time horizons:
          - Immediate  (24-72 hours) — critical risk reduction
          - Short-term (next sprint/patch cycle) — sustainable improvements
          - References — standards and documentation relevant to findings

        Format: labeled-group bullet lists with finding-type context tags.
        """
        doc          = self.doc
        remediations = getattr(self.enriched, 'remediations', []) or []

        doc.add_heading("Remediation Roadmap", level=1)
        T.add_horizontal_rule(doc)

        T.add_body_paragraph(
            doc,
            "The following remediation roadmap prioritizes identified findings into "
            "actionable time-bound phases. Immediate actions address the highest-severity "
            "risks. Short-term actions represent planned improvements for the next "
            "development or operations cycle."
        )
        T.add_spacer(doc)

        enriched_remediations = [r for r in remediations if getattr(r, 'enriched', False)]

        if not enriched_remediations:
            T.add_body_paragraph(
                doc,
                "Detailed AI-generated remediation guidance was not available for this session. "
                "Refer to the Technical Findings section for finding-specific recommendations, "
                "and consult relevant security standards (OWASP, CIS Benchmarks, NIST) for "
                "general remediation guidance."
            )
            return

        # ── Immediate actions (24-72 hours) ───────────────────────────────
        immediate_items = [
            (r.finding_type, action)
            for r in enriched_remediations
            for action in (getattr(r, 'immediate_actions', []) or [])
        ]

        if immediate_items:
            doc.add_heading("Immediate Actions  (24-72 Hours)", level=2)
            T.add_body_paragraph(
                doc,
                "The following actions should be taken immediately to reduce critical exposure:"
            )
            for finding_type, action in immediate_items:
                label = (finding_type or '').replace('_', ' ').title()
                para  = doc.add_paragraph(style='List Bullet')
                tag_run   = para.add_run(f"[{label}]  ")
                tag_run.bold = True
                tag_run.font.name = "Arial"
                tag_run.font.size = Pt(10.5)
                body_run  = para.add_run(action)
                body_run.font.name = "Arial"
                body_run.font.size = Pt(10.5)
                para.paragraph_format.space_after = Pt(3)
            T.add_spacer(doc)

        # ── Short-term actions ─────────────────────────────────────────────
        short_term_items = [
            (r.finding_type, action)
            for r in enriched_remediations
            for action in (getattr(r, 'short_term_actions', []) or [])
        ]

        if short_term_items:
            doc.add_heading("Short-Term Actions  (Next Sprint / Patch Cycle)", level=2)
            T.add_body_paragraph(
                doc,
                "The following improvements should be planned into the next development "
                "or operations sprint:"
            )
            for finding_type, action in short_term_items:
                label = (finding_type or '').replace('_', ' ').title()
                para  = doc.add_paragraph(style='List Bullet')
                tag_run   = para.add_run(f"[{label}]  ")
                tag_run.bold = True
                tag_run.font.name = "Arial"
                tag_run.font.size = Pt(10.5)
                body_run  = para.add_run(action)
                body_run.font.name = "Arial"
                body_run.font.size = Pt(10.5)
                para.paragraph_format.space_after = Pt(3)
            T.add_spacer(doc)

        # ── Technical commands ─────────────────────────────────────────────
        all_commands = [
            (r.finding_type, cmd)
            for r in enriched_remediations
            for cmd in (getattr(r, 'commands', []) or [])
        ]
        if all_commands:
            doc.add_heading("Technical Commands", level=2)
            T.add_body_paragraph(doc, "Copy-paste configuration commands for identified findings:")
            current_type = None
            for finding_type, cmd in all_commands:
                if finding_type != current_type:
                    label = (finding_type or '').replace('_', ' ').title()
                    T.add_label_paragraph(doc, label)
                    current_type = finding_type
                T.add_code_paragraph(doc, cmd)
            T.add_spacer(doc)

        # ── Standards and references ───────────────────────────────────────
        all_refs = sorted({
            ref
            for r in enriched_remediations
            for ref in (getattr(r, 'references', []) or [])
        })
        if all_refs:
            doc.add_heading("Standards & References", level=2)
            T.add_body_paragraph(doc, "Relevant security standards referenced in this assessment:")
            for ref in all_refs:
                T.add_bullet_paragraph(doc, ref)

    # =========================================================================
    # Section 8 — Appendix
    # =========================================================================

    def _build_appendix(self) -> None:
        """
        Build the Appendix section.

        Contents:
          - Session timeline (created, updated, file count)
          - Ingested file inventory from processing_log
          - Parser metadata (tool version, parse duration)
          - AI enrichment status and error log
          - Legal disclaimer
        """
        doc  = self.doc
        sess = self.session

        doc.add_heading("Appendix", level=1)
        T.add_horizontal_rule(doc)

        # ── A. Session timeline ────────────────────────────────────────────
        doc.add_heading("A.  Session Timeline", level=2)

        timeline_rows = [
            ("Session ID",       getattr(sess, 'session_id', '—')),
            ("Client Label",     getattr(sess, 'client_label', '—')),
            ("Status",           getattr(sess, 'session_status', '—')),
            ("Created",          _format_date(getattr(sess, 'created_at', ''))),
            ("Last Updated",     _format_date(getattr(sess, 'updated_at', ''))),
            ("Files Ingested",   str(getattr(sess, 'files_detected', 0))),
            ("Targets",          ", ".join(getattr(sess, 'targets', []) or []) or "—"),
            ("Tools",            ", ".join(getattr(sess, 'tools_detected', []) or []) or "—"),
        ]

        tl = doc.add_table(rows=len(timeline_rows), cols=2)
        T.apply_table_style(tl)
        for i, (k, v) in enumerate(timeline_rows):
            lc = tl.cell(i, 0)
            vc = tl.cell(i, 1)
            lc.width = Inches(2.0)
            vc.width = Inches(5.5)
            T.shade_cell(lc, T.COL_LIGHT_BLUE)
            T.set_cell_text(lc, k, bold=True, color_hex=T.COL_NAVY_HEX, font_size=10)
            T.set_cell_text(vc, v, font_size=10)

        T.add_spacer(doc)

        # ── B. Ingested file inventory ─────────────────────────────────────
        doc.add_heading("B.  Ingested Scan Files", level=2)

        processing_log = getattr(sess, 'processing_log', []) or []
        file_entries   = [e for e in processing_log if e.get("event") == "file_ingested"]

        if file_entries:
            hdr_cols   = ["Filename", "Tool", "Findings", "Parse (ms)", "Timestamp"]
            col_widths = [Inches(2.5), Inches(1.2), Inches(1.0), Inches(1.0), Inches(1.8)]

            fi = doc.add_table(rows=len(file_entries) + 1, cols=5)
            T.apply_table_style(fi)
            for j, (hdr, w) in enumerate(zip(hdr_cols, col_widths)):
                T.shade_cell(fi.cell(0, j), T.COL_NAVY_HEX)
                T.set_cell_text(fi.cell(0, j), hdr, bold=True, color_hex=T.COL_WHITE, font_size=9)
                fi.cell(0, j).width = w

            for i, entry in enumerate(file_entries, start=1):
                row_bg = T.COL_LIGHT_GREY if i % 2 == 0 else T.COL_WHITE
                for j in range(5):
                    T.shade_cell(fi.cell(i, j), row_bg)
                ts = entry.get("timestamp", "—")
                if ts and len(ts) > 19:
                    ts = ts[:19].replace("T", " ")
                T.set_cell_text(fi.cell(i, 0), entry.get("filename", "—"),              font_size=9)
                T.set_cell_text(fi.cell(i, 1), entry.get("tool_type", "—"),             font_size=9)
                T.set_cell_text(fi.cell(i, 2), str(entry.get("parse_findings", 0)),     font_size=9)
                T.set_cell_text(fi.cell(i, 3), str(entry.get("parse_ms", "—")),         font_size=9)
                T.set_cell_text(fi.cell(i, 4), ts,                                       font_size=9)
        else:
            T.add_body_paragraph(doc, "No file ingestion entries were found in the session log.")

        T.add_spacer(doc)

        # ── C. Parser metadata ─────────────────────────────────────────────
        doc.add_heading("C.  Parser Metadata", level=2)
        if self.parsed_data_list:
            for parsed in self.parsed_data_list:
                tool_type = getattr(parsed, 'tool_type', 'UNKNOWN')
                version   = getattr(parsed, 'tool_version', None) or 'version unknown'
                duration  = getattr(parsed, 'parse_duration_ms', 0)
                T.add_body_paragraph(
                    doc,
                    f"{tool_type}: {version}  (parsed in {duration:.1f} ms)"
                )
        else:
            T.add_body_paragraph(doc, "No parser metadata available.")

        T.add_spacer(doc)

        # ── D. AI enrichment status ────────────────────────────────────────
        doc.add_heading("D.  AI Enrichment Status", level=2)
        is_complete = getattr(self.enriched, 'enrichment_complete', False)
        status_str  = "Complete" if is_complete else "Partial / Unavailable"
        T.add_body_paragraph(doc, f"Enrichment status: {status_str}")
        T.add_body_paragraph(
            doc,
            f"Findings enriched: {self.enriched.finding_count}"
        )

        enrich_errors = getattr(self.enriched, 'enrichment_errors', []) or []
        if enrich_errors:
            T.add_label_paragraph(doc, "Enrichment Errors")
            for err in enrich_errors:
                T.add_bullet_paragraph(doc, err)

        T.add_spacer(doc)

        # ── E. Disclaimer ──────────────────────────────────────────────────
        doc.add_heading("E.  Disclaimer", level=2)
        T.add_body_paragraph(
            doc,
            "This report was generated by IntelliAssess AI based on non-intrusive external "
            "reconnaissance techniques. All findings reflect observations at the time of "
            "assessment and may not account for changes made to the assessed environment "
            "after the assessment date."
        )
        T.add_body_paragraph(
            doc,
            "This report does not constitute a guarantee of security, a comprehensive "
            "penetration test, or a formal compliance audit. Compliance determinations "
            "require review by qualified auditors and legal counsel. IntelliAssess AI "
            "findings are indicators for further investigation, not conclusive proof of "
            "exploitability."
        )

    # =========================================================================
    # Private helpers
    # =========================================================================

    def _compute_severity_counts(self) -> dict[str, int]:
        """Return a severity → count dict from all finding_summaries."""
        counts: dict[str, int] = {}
        for f in self.enriched.finding_summaries:
            sev = (getattr(f, 'severity_label', 'INFO') or 'INFO').upper()
            counts[sev] = counts.get(sev, 0) + 1
        return counts

    def _build_severity_summary_table(self, severity_counts: dict[str, int]) -> None:
        """
        Render a color-coded severity count summary table.

        Severities are shown in priority order: CRITICAL → HIGH → MEDIUM → LOW → INFO.
        """
        doc = self.doc
        ordered = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"]
        rows    = [(sev, severity_counts[sev]) for sev in ordered if sev in severity_counts]

        if not rows:
            return

        sev_table = doc.add_table(rows=len(rows) + 1, cols=2)
        T.apply_table_style(sev_table)

        T.shade_cell(sev_table.cell(0, 0), T.COL_NAVY_HEX)
        T.shade_cell(sev_table.cell(0, 1), T.COL_NAVY_HEX)
        T.set_cell_text(sev_table.cell(0, 0), "Severity", bold=True, color_hex=T.COL_WHITE, font_size=10)
        T.set_cell_text(sev_table.cell(0, 1), "Count",    bold=True, color_hex=T.COL_WHITE, font_size=10)
        sev_table.cell(0, 0).width = Inches(3.0)
        sev_table.cell(0, 1).width = Inches(4.5)

        for i, (sev, count) in enumerate(rows, start=1):
            bg_hex, fg_hex = T.get_severity_colors(sev)
            row_bg = T.COL_LIGHT_GREY if i % 2 == 0 else T.COL_WHITE
            T.shade_cell(sev_table.cell(i, 0), bg_hex)
            T.shade_cell(sev_table.cell(i, 1), row_bg)
            T.set_cell_text(sev_table.cell(i, 0), sev,       bold=True, color_hex=fg_hex)
            T.set_cell_text(sev_table.cell(i, 1), str(count))

    def _collect_weak_cipher_names(self, target: str) -> set[str]:
        """
        Return a set of weak cipher suite names for a given target.

        Cross-references WEAK_CIPHER findings from finding_summaries.
        Used by _render_ssl_block to annotate the cipher inventory display.
        """
        weak_ciphers: set[str] = set()
        for f in self.enriched.finding_summaries:
            if (
                getattr(f, 'finding_type', '') == 'WEAK_CIPHER'
                and getattr(f, 'target', '') == target
            ):
                detail = getattr(f, 'raw_finding_detail', '') or ''
                # raw_finding_detail format: "Weak cipher accepted: CIPHER_NAME"
                # Extract the cipher name from the detail string
                if ':' in detail:
                    candidate = detail.split(':', 1)[-1].strip()
                    if candidate:
                        weak_ciphers.add(candidate)
        return weak_ciphers


# ---------------------------------------------------------------------------
# Module-level utilities
# ---------------------------------------------------------------------------

def _format_date(iso_str: str) -> str:
    """
    Format an ISO 8601 timestamp string to a readable date string.

    Returns 'YYYY-MM-DD HH:MM UTC' or the original string if parsing fails.
    Falls back to '—' for empty input.
    """
    if not iso_str:
        return "—"
    try:
        dt = datetime.fromisoformat(iso_str.replace("Z", "+00:00"))
        return dt.strftime("%Y-%m-%d  %H:%M UTC")
    except (ValueError, AttributeError):
        # Return the raw string if it's not parseable — better than losing it
        return iso_str[:19].replace("T", " ") if len(iso_str) >= 19 else iso_str


def _ctx_value(ctx: dict, key: str) -> str:
    """Return a context dict value or '—' if missing or None."""
    val = ctx.get(key)
    return str(val).title() if val else "—"
