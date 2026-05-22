"""
reporting/templates.py
=======================
Pure formatting layer for IntelliAssess AI DOCX report generation — Phase 4-2.

Responsibility: formatting primitives ONLY.
  - Color palette constants
  - Document style configuration (fonts, headings, margins)
  - Paragraph / bullet / label helper functions
  - Table border and cell shading utilities
  - Horizontal rule / spacer helpers

Design principles:
  - Zero business logic here — this module knows nothing about findings,
    sessions, or enrichment data.
  - All OOXML manipulation is isolated in this module.
  - reporter.py calls these helpers; it never touches XML directly.
  - All helpers are defensive: they never raise on bad input.

Used exclusively by: reporting/reporter.py
"""

from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Color palette — hex strings for OOXML, RGB tuples for python-docx runs
# ---------------------------------------------------------------------------

# Brand / structural colors
COL_NAVY_HEX   = "1B2A47"   # Primary: H1, cover background accent, table headers
COL_BLUE_HEX   = "2E75B6"   # Secondary: H2, label text, accent rule color
COL_LIGHT_BLUE = "D5E8F0"   # Subtle label-column shading in metadata tables
COL_GREY_HEX   = "4A4A4A"   # Default body text (off-black for readability)
COL_LIGHT_GREY = "F5F5F5"   # Alternating row background
COL_WHITE      = "FFFFFF"   # White cell / text

# Severity badge palette — (background hex, foreground text hex)
# Used for finding severity cells and risk posture indicator.
SEVERITY_COLORS: dict[str, tuple[str, str]] = {
    "CRITICAL": ("B80000", "FFFFFF"),   # Deep red
    "HIGH":     ("D04800", "FFFFFF"),   # Orange-red
    "MEDIUM":   ("C8890A", "FFFFFF"),   # Amber-brown (readable on white)
    "LOW":      ("2E75B6", "FFFFFF"),   # Blue
    "INFO":     ("357A38", "FFFFFF"),   # Forest green
    "UNKNOWN":  ("787878", "FFFFFF"),   # Neutral grey
}


# ---------------------------------------------------------------------------
# Document style configuration
# ---------------------------------------------------------------------------

def configure_document_styles(doc: Document) -> None:
    """
    Apply professional style defaults to a freshly created Document.

    Sets:
      - Page size: US Letter (8.5 × 11 inches), 1-inch margins
      - Default body font: Arial 10.5 pt
      - Heading 1: Arial 16 pt bold navy — primary section headings
      - Heading 2: Arial 13 pt bold blue — subsection headings
      - Heading 3: Arial 11 pt bold grey — labeled subsections

    Called once by DocxReporter.__init__() before any section is built.
    """
    # ── Page size and margins (US Letter) ─────────────────────────────────
    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Default body font ──────────────────────────────────────────────────
    normal = doc.styles['Normal']
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    # ── Heading 1 — navy ──────────────────────────────────────────────────
    h1 = doc.styles['Heading 1']
    h1.font.name  = "Arial"
    h1.font.size  = Pt(16)
    h1.font.bold  = True
    h1.font.color.rgb = RGBColor(0x1B, 0x2A, 0x47)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(4)

    # ── Heading 2 — blue ──────────────────────────────────────────────────
    h2 = doc.styles['Heading 2']
    h2.font.name  = "Arial"
    h2.font.size  = Pt(13)
    h2.font.bold  = True
    h2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after  = Pt(3)

    # ── Heading 3 — dark grey ─────────────────────────────────────────────
    h3 = doc.styles['Heading 3']
    h3.font.name  = "Arial"
    h3.font.size  = Pt(11)
    h3.font.bold  = True
    h3.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    h3.paragraph_format.space_before = Pt(7)
    h3.paragraph_format.space_after  = Pt(2)


# ---------------------------------------------------------------------------
# Paragraph helpers
# ---------------------------------------------------------------------------

def add_body_paragraph(
    doc:       Document,
    text:      str,
    bold:      bool = False,
    color_hex: str  = COL_GREY_HEX,
) -> None:
    """
    Add a normal body paragraph with Arial 10.5 pt styling.

    Args:
        doc:       Target document.
        text:      Paragraph text.
        bold:      Whether the run is bold.
        color_hex: Hex color for the text (default: dark grey).
    """
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.bold = bold
    r, g, b = hex_to_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    para.paragraph_format.space_after = Pt(6)


def add_label_paragraph(doc: Document, text: str) -> None:
    """
    Add a small uppercase label paragraph — acts as an inline subsection marker.

    Style: Arial 9 pt, bold, blue, small top spacing.
    Used above body paragraphs to label their semantic role
    (e.g. "OBSERVATION", "BUSINESS IMPACT", "RECOMMENDATION").
    """
    para = doc.add_paragraph()
    run  = para.add_run(text.upper())
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    r, g, b = hex_to_rgb(COL_BLUE_HEX)
    run.font.color.rgb = RGBColor(r, g, b)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(1)


def add_bullet_paragraph(doc: Document, text: str) -> None:
    """
    Add a bullet list item using the built-in 'List Bullet' style.

    Never manually inserts bullet characters — uses docx list style.
    """
    para = doc.add_paragraph(style='List Bullet')
    run  = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    para.paragraph_format.space_after = Pt(2)


def add_code_paragraph(doc: Document, text: str) -> None:
    """
    Add a monospaced code/command paragraph with a light grey background.

    Used for technical commands in the Remediation section.
    """
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    # Light background via paragraph shading
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'),  "EFEFEF")
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'),   'clear')
    pPr.append(shd)
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)


def add_horizontal_rule(doc: Document) -> None:
    """
    Add a thin horizontal rule using a bottom-border on an empty paragraph.

    Color matches the blue accent. Used after section H1 headings.
    """
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), COL_BLUE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(8)


def add_spacer(doc: Document, pt: int = 6) -> None:
    """Add an empty paragraph as a vertical spacer with given point height."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(pt)


# ---------------------------------------------------------------------------
# Table cell helpers
# ---------------------------------------------------------------------------

def shade_cell(cell, hex_color: str) -> None:
    """
    Apply a solid background fill to a table cell via OOXML.

    Args:
        cell:      A python-docx TableCell object.
        hex_color: 6-character hex color string (no leading #).
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'),  hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'),   'clear')
    tcPr.append(shd)


def set_cell_text(
    cell,
    text:       str,
    bold:       bool  = False,
    color_hex:  str   = COL_GREY_HEX,
    font_size:  float = 10.5,
    center:     bool  = False,
    top_pad:    int   = 60,
    bottom_pad: int   = 60,
) -> None:
    """
    Write styled text into a table cell.

    Clears any existing cell content before writing. Adds internal cell
    padding for visual breathing room.

    Args:
        cell:       Target TableCell.
        text:       Cell text content.
        bold:       Whether to bold the run.
        color_hex:  Text color hex (6-char, no #).
        font_size:  Font size in points.
        center:     Whether to centre-align the paragraph.
        top_pad:    Top internal padding in twentieths of a point (DXA-ish).
        bottom_pad: Bottom internal padding.
    """
    cell.text = ""
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    r, g, b = hex_to_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)

    # Internal cell padding via tcMar
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top_pad), ('bottom', bottom_pad), ('left', 120), ('right', 120)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'),    str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def apply_table_style(table) -> None:
    """
    Apply a clean thin-border style to a table.

    Replaces default docx table grid with subtle light-grey borders
    on all sides and internal dividers.
    """
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove any existing tblBorders node to avoid duplicates
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# Severity utilities
# ---------------------------------------------------------------------------

def get_severity_colors(severity: str) -> tuple[str, str]:
    """
    Return (background_hex, foreground_hex) for a severity level.

    Falls back to UNKNOWN grey if the severity is not recognized.
    """
    return SEVERITY_COLORS.get(severity.upper(), SEVERITY_COLORS["UNKNOWN"])


# ---------------------------------------------------------------------------
# Color conversion utility
# ---------------------------------------------------------------------------

def hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    """
    Convert a 6-character hex color string to an (r, g, b) integer tuple.

    Args:
        hex_color: Hex string with or without leading '#'.

    Returns:
        (r, g, b) tuple with values 0-255.

    Example:
        hex_to_rgb("1B2A47") → (27, 42, 71)
    """
    h = hex_color.lstrip('#')
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
